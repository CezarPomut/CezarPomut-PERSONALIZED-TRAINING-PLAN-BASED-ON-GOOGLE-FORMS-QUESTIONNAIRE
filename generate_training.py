'''facem o functie care sa aiba un parametru de intrare un ob, ne va genera un fisier word sau txt, care sa aiba numele persoanei,
si in interior sa scriem planul de antrenament'''
from docx import *
from select_training import *

def generare_antrenament(obj):
    doc = Document()
    doc.add_paragraph(obj.get_nume())
    counter = 1
    for lista1 in obj.antrenament:
        ziua="Antrenament ziua "+str(counter)
        doc.add_paragraph(ziua)
        counter=counter+1
        for string in lista1:
            doc.add_paragraph("     "+string)
    path="C:/Users/pomut/Desktop/"+obj.get_nume()+".docx"
    doc.save(path)




